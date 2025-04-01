from http.server import BaseHTTPRequestHandler, HTTPServer
import json
from docx import Document
from io import BytesIO

class WordHandler(BaseHTTPRequestHandler):
    def do_POST(self):
        if self.path == '/generate-word':
            content_length = int(self.headers['Content-Length'])
            body = self.rfile.read(content_length)
            data = json.loads(body)

            # Load template
            doc = Document('allergist-prescription-template.docx')
            paragraphs = [p.text for p in doc.paragraphs]
            print("Loaded Paragraphs:", paragraphs)  # Debugging output

            # Fill template
            replacements = {
                'Name:': f'Name: {data["name"]}',
                'Age:': f'Age: {data["age"]}',
                'Sex:': f'Sex: {data["sex"]}',
                'Diagnosis:': f'Diagnosis: {data["diagnosis"]}',
                'Date:': f'Date: {data["date"]}'
            }

            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)

            # Save to memory
            output = BytesIO()
            doc.save(output)
            output.seek(0)

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.send_header('Content-Disposition', 'attachment; filename=filled-form.docx')
            self.end_headers()
            self.wfile.write(output.read())

if __name__ == '__main__':
    server_address = ('', 8000)
    httpd = HTTPServer(server_address, WordHandler)
    print('Server running on http://localhost:8000')
    httpd.serve_forever()
